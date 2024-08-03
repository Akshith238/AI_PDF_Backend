from flask import Flask, request, jsonify
from flask_cors import CORS
import os
import torch
import re
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from faster_whisper import WhisperModel
from concurrent.futures import ThreadPoolExecutor
import time
from pathlib import Path as p
from pprint import pprint
from langchain_google_genai import ChatGoogleGenerativeAI
import pandas as pd
from langchain import PromptTemplate
from langchain.chains.question_answering import load_qa_chain
from langchain.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Chroma
from langchain.chains import RetrievalQA
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
import base64
import PIL.Image
import matplotlib.pyplot as plt
from PyPDF2 import PdfWriter, PdfReader
import PyPDF2
from reportlab.lib.utils import ImageReader
from PIL import Image
from io import BytesIO
import pathlib
import shutil
from new import func
import mysql.connector
from docx import Document
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"

app = Flask(__name__)
CORS(app)  # Enable CORS for all routes

# Define model size and device (adjust as needed)
model_size = "distil-large-v3"
device = "cuda" if torch.cuda.is_available() else "cpu"  # Use GPU if available

# Load the Whisper model
whisper_model = WhisperModel(model_size, device=device)

# Define paths for file uploads and PDF output
UPLOAD_FOLDER = 'uploads'
PDF_FOLDER = 'pdfs'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(PDF_FOLDER, exist_ok=True)

# Words to remove pattern (case-insensitive)
words_to_remove_pattern = r"\b(yep|um|so|okay|sorry\|yes|yeah\.|ok\.|thank you\.|sure\.|please|and|oh,|no,|you |good, yeah.|I'm sorry.)\b"
compiled_pattern = re.compile(words_to_remove_pattern, flags=re.IGNORECASE)

executor = ThreadPoolExecutor(max_workers=2)  # Adjust the number of workers as needed
db_config = {
    'host': 'localhost',
    'database': 'caterpillar',
    'user': 'root',
    'password': 'root'
}


def connect_to_database():
    return mysql.connector.connect(**db_config)


@app.route('/upload', methods=['POST'])
def upload_files():
    try:
        # Define the upload folder
        UPLOAD_FOLDER = 'uploads'  # Replace with your actual upload folder path

        # Function to clear the upload folder
        def clear_upload_folder(folder):
            for filename in os.listdir(folder):
                file_path = os.path.join(folder, filename)
                try:
                    if os.path.isfile(file_path) or os.path.islink(file_path):
                        os.unlink(file_path)
                    elif os.path.isdir(file_path):
                        shutil.rmtree(file_path)
                except Exception as e:
                    print(f'Failed to delete {file_path}. Reason: {e}')
        
        # Clear the upload folder
        clear_upload_folder(UPLOAD_FOLDER)

        # Handle images
        image_paths = []
        i = 0
        while True:
            key = f'image{i}'
            if key in request.files:
                file = request.files[key]
                filename = f'image{i + 1}.jpg'  # Adjust filename as needed
                file.save(os.path.join(UPLOAD_FOLDER, filename))
                image_paths.append(os.path.join(UPLOAD_FOLDER, filename))
                i += 1
            else:
                break

        # Handle audio
        if 'audio' in request.files:
            audio_file = request.files['audio']
            audio_filename = 'audio.wav'
            audio_path = os.path.join(UPLOAD_FOLDER, audio_filename)
            audio_file.save(audio_path)
        else:
            audio_path = None

        # Run PDF generation in the background
        future = executor.submit(generate_pdf_main, audio_path, image_paths)
        future.add_done_callback(lambda x: print(f"PDF generation result: {x.result()}"))

        # Return success message with PDF path
        return jsonify({'message': 'Files uploaded successfully, PDF generation in progress.'}), 200

    except Exception as e:
        return jsonify({'error': f'Error uploading files: {str(e)}'}), 500


def generate_pdf_main(audio_path, image_paths):
    import os
    import warnings
    from pathlib import Path as p
    from pprint import pprint
    from langchain_google_genai import ChatGoogleGenerativeAI
    import pandas as pd
    from langchain import PromptTemplate
    from langchain.chains.question_answering import load_qa_chain
    from langchain.document_loaders import PyPDFLoader
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain.vectorstores import Chroma
    from langchain.chains import RetrievalQA
    from langchain_google_genai import GoogleGenerativeAIEmbeddings
    import google.generativeai as genai
    import base64
    import PIL.Image
    import matplotlib.pyplot as plt
    import google.generativeai as genai
    from PyPDF2 import PdfWriter, PdfReader
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.utils import ImageReader
    from PIL import Image
    from io import BytesIO
    import pathlib
    warnings.filterwarnings("ignore")
    from docx import Document
    import PyPDF2
    

    try:
        if audio_path:
            # Perform transcription
            segments, info = whisper_model.transcribe(audio_path, beam_size=1)
        else:
            segments = []

        # Prepare PDF
        pdf_filename = os.path.join(PDF_FOLDER, f'transcript.pdf')
        c = canvas.Canvas(pdf_filename, pagesize=letter)
        width, height = letter
        text = c.beginText(40, height - 40)
        text.setFont("Helvetica", 12)

        # Process and store results
        for segment in segments:
            filtered_text = compiled_pattern.sub("", segment.text)
            filtered_text = re.sub('[^A-Za-z0-9]+', ' ', filtered_text).strip()
            text.textLine(filtered_text)

        c.drawText(text)
        c.save()
        genai.configure(api_key='AIzaSyCqEKwd23ztVuk-dkCXypjeHWlcs41aCSM') 
        import os


        # Function to extract text from PDF
        def extract_text_from_pdf(pdf_path):
            with open(pdf_path, "rb") as f:
                pdf_reader = PyPDF2.PdfReader(f)
                text = ""
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text()
                return text

        # Function to extract text from DOCX
        def extract_text_from_docx(docx_path):
            doc = Document(docx_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return '\n\n'.join(full_text)

        # Function to extract text from TXT
        def extract_text_from_txt(txt_path):
            with open(txt_path, "r", encoding='utf-8') as f:
                return f.read()

        # List of paths to your input files (PDF, DOCX, TXT)
        file_paths = ["D:/PSG tech/Caterpillar/Caterpillar-ChillyFlakes/backend/pdfs/transcript.pdf"]

        # Extract text from files
        texts = []
        for path in file_paths:
            if path.endswith(".pdf"):
                texts.append(extract_text_from_pdf(path))
            elif path.endswith(".docx"):
                texts.append(extract_text_from_docx(path))
            elif path.endswith(".txt"):
                texts.append(extract_text_from_txt(path))

        context = "\n\n".join(texts)

        # Initialize text splitter and embeddings
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
        texts = text_splitter.split_text(context)


        model = ChatGoogleGenerativeAI(model="gemini-pro",google_api_key="AIzaSyCqEKwd23ztVuk-dkCXypjeHWlcs41aCSM",
                                    temperature=0.2,convert_system_message_to_human=True)
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key="AIzaSyCqEKwd23ztVuk-dkCXypjeHWlcs41aCSM")
        vector_index = Chroma.from_texts(texts, embeddings).as_retriever(search_kwargs={"k":5})

        qa_chain = RetrievalQA.from_chain_type(
            model,
            retriever=vector_index,
            return_source_documents=True
        )
        template = """
        You are the Assistant to a Truck Inspector Based On the context you provide a report which the user easily understands and structured and highlight the problems If the condition is - Needs Replacement, Bad ,Broken, High, Low, Rust give a detailed explanation about it answer only from the context\n\n
        Give the report summary so that the customer understands the condition of the truck\n\n
        {context}
        Question: {question}
        Helpful Answer:"""
        QA_CHAIN_PROMPT = PromptTemplate.from_template(template)
        qa_chain = RetrievalQA.from_chain_type(
            model,
            retriever=vector_index,
            return_source_documents=True,
            chain_type_kwargs={"prompt": QA_CHAIN_PROMPT}
        )

        question = """
        Generate a brief report on the truck condition 
        """
        result = qa_chain({"query": question})

        formatreport=str(result["result"])
        print(formatreport)     

        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        def generate_pdf(answer, file_name="report.pdf"):
            c = canvas.Canvas(file_name, pagesize=letter)
            width, height = letter

            c.setFont("Helvetica-Bold", 14)
            c.drawString(72, height - 72, "Vehicle Condition Report")

            c.setFont("Helvetica", 12)
            text = c.beginText(72, height - 100)
            text.setTextOrigin(72, height - 100)
            text.setLeading(14.5)

            lines = answer.split('\n')
            for line in lines:
                text.textLine(line)

            c.drawText(text)
            c.save()
            print("saved report condition")

        generate_pdf(formatreport, "condition_report.pdf")



        file_paths = ["D:/PSG tech/Caterpillar/Caterpillar-ChillyFlakes/backend/condition_report.pdf","D:/PSG tech/Caterpillar/Caterpillar-ChillyFlakes/backend/data/correct data.docx"]

        texts = []
        for path in file_paths:
            if path.endswith(".pdf"):
                texts.append(extract_text_from_pdf(path))
            elif path.endswith(".docx"):
                texts.append(extract_text_from_docx(path))
            elif path.endswith(".txt"):
                texts.append(extract_text_from_txt(path))

        context = "\n\n".join(texts)

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
        texts = text_splitter.split_text(context)

        template = """
        Based on the detailes of the truck generate a complete and brief report highlighting the issues found and provide suggestion for it by making comparisons with the correct value and provide suggestions only for that\n\n
        {context}
        Question: {question}
        Helpful Answer:"""
        QA_CHAIN_PROMPT = PromptTemplate.from_template(template)
        qa_chain = RetrievalQA.from_chain_type(
            model,
            retriever=vector_index,
            return_source_documents=True,
            chain_type_kwargs={"prompt": QA_CHAIN_PROMPT}
        )

        question = """
        state the issue and the suggestion and state the area where there is a issue and compare it with the correct data  """
        result = qa_chain({"query": question})
        result["result"]
        ans=str(result["result"])
        print(ans)

        finalreport=formatreport
        finalreport+="\n\n"
        finalreport+=ans
        print("\n\n",finalreport,"final report \n\n")

        from fpdf import FPDF


        pdf = FPDF()

        pdf.add_page()

        pdf.set_font("Arial", size = 10)

        pdf.multi_cell(0, 4, finalreport)

        pdf_output = "finalreport.pdf"
        pdf.output(pdf_output)

        print(f"PDF successfully created and saved as {pdf_output}")

        ################################image 



        def add_image_to_pdf(pdf_path, image_path, output_path):
            img = Image.open(image_path)
            orig_width, orig_height = img.size

            max_size = min(orig_width, orig_height)
            new_width = new_height = min(max_size, 150) 

            img = img.resize((new_width, new_height), Image.LANCZOS)


            packet = BytesIO()
            can = canvas.Canvas(packet, pagesize=letter)

            text_margin = 10  
            genai.configure(api_key='AIzaSyCqEKwd23ztVuk-dkCXypjeHWlcs41aCSM')
            model = genai.GenerativeModel('gemini-1.5-flash')

            image1 = {
            'mime_type': 'image/jpeg',
            'data': pathlib.Path(image_path).read_bytes()
            }

            prompt = "One title of the image?"

            response = model.generate_content([prompt, image1])
            print(response.text)

            can.drawString(100, 700, response.text)

            img_y = 700 - text_margin - new_height  

            # Draw the image on the canvas
            can.drawImage(ImageReader(img), 50, img_y, width=new_width, height=new_height)
            can.showPage()
            can.save()

            packet.seek(0)

            existing_pdf = PdfReader(open(pdf_path, "rb"))
            output = PdfWriter()

            for page in existing_pdf.pages:
                output.add_page(page)

            new_pdf = PdfReader(packet)
            output.add_page(new_pdf.pages[0])

            with open(output_path, "wb") as output_stream:
                output.write(output_stream)



        image_directory = "uploads"

        final_pdf_path = "finalreport.pdf"


        for image_file in os.listdir(image_directory):
            if image_file.endswith(".jpg") or image_file.endswith(".jpeg") or image_file.endswith(".png"):
                image_path = os.path.join(image_directory, image_file)

                add_image_to_pdf(final_pdf_path, image_path, final_pdf_path)

        print("PDF generation complete.")




        file_paths = ["finalreport.pdf","data/Lists.docx"]


        # Function to extract text from PDF
        def extract_text_from_pdf(pdf_path):
            with open(pdf_path, "rb") as f:
                pdf_reader = PyPDF2.PdfReader(f)
                text = ""
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text()
                return text

        # Function to extract text from DOCX
        def extract_text_from_docx(docx_path):
            doc = Document(docx_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return '\n\n'.join(full_text)

        # Function to extract text from TXT
        def extract_text_from_txt(txt_path):
            with open(txt_path, "r", encoding='utf-8') as f:
                return f.read()
        texts = []
        for path in file_paths:
            if path.endswith(".pdf"):
                texts.append(extract_text_from_pdf(path))
            elif path.endswith(".docx"):
                texts.append(extract_text_from_docx(path))
            elif path.endswith(".txt"):
                texts.append(extract_text_from_txt(path))

        context = "\n\n".join(texts)

        # Initialize text splitter and embeddings
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
        texts = text_splitter.split_text(context)
        model = ChatGoogleGenerativeAI(model="gemini-pro",google_api_key="AIzaSyCqEKwd23ztVuk-dkCXypjeHWlcs41aCSM",
                                    temperature=0.2,convert_system_message_to_human=True)
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key="AIzaSyCqEKwd23ztVuk-dkCXypjeHWlcs41aCSM")
        vector_index = Chroma.from_texts(texts, embeddings).as_retriever(search_kwargs={"k":5})

        qa_chain = RetrievalQA.from_chain_type(
            model,
            retriever=vector_index,
            return_source_documents=True
        )


        template = """
        analyse the context \n\n
        List down the sections and sub sections that are metioned in the context and compare with the section needed to be inspected and list the missing \n
        Example:
        user:
        * Engine Oil Condition: Good
            * Engine Oil Color: Clean
            * Oil Leak on Engine: No

            Sections Needed to be Inspected:
            *Engine:*
            * Engine Oil Color
            * Oil leak on engine

        system:
        None
        {context}
        Question: {question}
        Helpful Answer:"""
        QA_CHAIN_PROMPT = PromptTemplate.from_template(template)
        qa_chain = RetrievalQA.from_chain_type(
            model,
            retriever=vector_index,
            return_source_documents=True,
            chain_type_kwargs={"prompt": QA_CHAIN_PROMPT}
        )

        question = """
        List down only the missing inspectable items
        """
        result = qa_chain({"query": question})
        ans=str(result["result"])
        print(ans)
        func()
        return 'finalreport.pdf'

    except Exception as e:
        raise RuntimeError(f'Error generating PDF: {str(e)}')

@app.route('/reports', methods=['GET'])
def get_reports():
    try:
        # Connect to the database
        connection = connect_to_database()

        # Prepare a cursor object using cursor() method
        cursor = connection.cursor()

        # Query to fetch all reports
        query = "SELECT id, pdf_data, repstatus FROM pdfs"
        cursor.execute(query)

        # Fetch all rows
        reports = cursor.fetchall()

        # Close cursor and connection
        cursor.close()
        connection.close()

        # Convert to a list of dictionaries for JSON response
        reports_list = []
        for report in reports:
            report_dict = {
                'id': report[0],
                'repstatus': report[2],
                # Convert pdf_data to a base64 encoded string
                'pdf_data': report_to_base64(report[1]),
                # Add more fields as needed
            }
            reports_list.append(report_dict)

        # Return JSON response
        return jsonify({'reports': reports_list}), 200

    except mysql.connector.Error as e:
        return jsonify({'error': f'Failed to fetch reports: {str(e)}'}), 500

    except Exception as e:
        return jsonify({'error': f'Unknown error occurred: {str(e)}'}), 500

def report_to_base64(pdf_data):
    # Convert bytes to base64 encoded string
    return base64.b64encode(pdf_data).decode('utf-8')
        
        
if __name__ == '__main__':
    app.run(debug=True)
