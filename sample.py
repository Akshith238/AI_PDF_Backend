import warnings
from pathlib import Path as p
from pprint import pprint
from langchain_google_genai import ChatGoogleGenerativeAI
import pandas as pd
from langchain import PromptTemplate
from langchain.chains.question_answering import load_qa_chain
from langchain.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Chroma
from langchain.chains import RetrievalQA
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
import base64
import PIL.Image
import matplotlib.pyplot as plt
import google.generativeai as genai
from PyPDF2 import PdfWriter, PdfReader
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader
from PIL import Image
from io import BytesIO
import pathlib
warnings.filterwarnings("ignore")
from docx import Document
import PyPDF2

genai.configure(api_key='AIzaSyCqEKwd23ztVuk-dkCXypjeHWlcs41aCSM') 
import os


# Function to extract text from PDF
def extract_text_from_pdf(pdf_path):
    with open(pdf_path, "rb") as f:
        pdf_reader = PyPDF2.PdfReader(f)
        text = ""
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text()
        return text

# Function to extract text from DOCX
def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n\n'.join(full_text)

# Function to extract text from TXT
def extract_text_from_txt(txt_path):
    with open(txt_path, "r", encoding='utf-8') as f:
        return f.read()

# List of paths to your input files (PDF, DOCX, TXT)
file_paths = ["D:/PSG tech/Caterpillar/Caterpillar-ChillyFlakes/backend/pdfs/transcript.pdf"]

# Extract text from files
texts = []
for path in file_paths:
    if path.endswith(".pdf"):
        texts.append(extract_text_from_pdf(path))
    elif path.endswith(".docx"):
        texts.append(extract_text_from_docx(path))
    elif path.endswith(".txt"):
        texts.append(extract_text_from_txt(path))

context = "\n\n".join(texts)

# Initialize text splitter and embeddings
text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
texts = text_splitter.split_text(context)


model = ChatGoogleGenerativeAI(model="gemini-pro",google_api_key="AIzaSyCqEKwd23ztVuk-dkCXypjeHWlcs41aCSM",temperature=0.2,convert_system_message_to_human=True)
embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key="AIzaSyCqEKwd23ztVuk-dkCXypjeHWlcs41aCSM")
vector_index = Chroma.from_texts(texts, embeddings).as_retriever(search_kwargs={"k":5})

qa_chain = RetrievalQA.from_chain_type(
    model,
    retriever=vector_index,
    return_source_documents=True
)
template = """
You are the Assistant to a Truck Inspector Based On the context you provide a report which the user easily understands and structured and highlight the problems If the condition is - Needs Replacement, Bad ,Broken, High, Low, Rust give a detailed explanation about it answer only from the context\n\n
Give the report summary so that the customer understands the condition of the truck\n\n
{context}
Question: {question}
Helpful Answer:"""
QA_CHAIN_PROMPT = PromptTemplate.from_template(template)
qa_chain = RetrievalQA.from_chain_type(
    model,
    retriever=vector_index,
    return_source_documents=True,
    chain_type_kwargs={"prompt": QA_CHAIN_PROMPT}
)

question = """
Generate a brief report on the truck condition 
"""
result = qa_chain({"query": question})

formatreport=str(result["result"])
print(formatreport)     

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

def generate_pdf(answer, file_name="report.pdf"):
    c = canvas.Canvas(file_name, pagesize=letter)
    width, height = letter

    c.setFont("Helvetica-Bold", 14)
    c.drawString(72, height - 72, "Vehicle Condition Report")

    c.setFont("Helvetica", 12)
    text = c.beginText(72, height - 100)
    text.setTextOrigin(72, height - 100)
    text.setLeading(14.5)

    lines = answer.split('\n')
    for line in lines:
        text.textLine(line)

    c.drawText(text)
    c.save()
    print("saved report condition")

generate_pdf(formatreport, "condition_report.pdf")



file_paths = ["D:/PSG tech/Caterpillar/Caterpillar-ChillyFlakes/backend/condition_report.pdf","D:/PSG tech/Caterpillar/Caterpillar-ChillyFlakes/backend/data/correct data.docx"]

texts = []
for path in file_paths:
    if path.endswith(".pdf"):
        texts.append(extract_text_from_pdf(path))
    elif path.endswith(".docx"):
        texts.append(extract_text_from_docx(path))
    elif path.endswith(".txt"):
        texts.append(extract_text_from_txt(path))

context = "\n\n".join(texts)

text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
texts = text_splitter.split_text(context)

template = """
Based on the detailes of the truck generate a complete and brief report highlighting the issues found and provide suggestion for it by making comparisons with the correct value and provide suggestions only for that\n\n
{context}
Question: {question}
Helpful Answer:"""
QA_CHAIN_PROMPT = PromptTemplate.from_template(template)
qa_chain = RetrievalQA.from_chain_type(
    model,
    retriever=vector_index,
    return_source_documents=True,
    chain_type_kwargs={"prompt": QA_CHAIN_PROMPT}
)

question = """
state the issue and the suggestion and state the area where there is a issue and compare it with the correct data  """
result = qa_chain({"query": question})
result["result"]
ans=str(result["result"])
print(ans)

finalreport=formatreport
finalreport+="\n\n"
finalreport+=ans
print("\n\n",finalreport,"final report \n\n")

from fpdf import FPDF


pdf = FPDF()

pdf.add_page()

pdf.set_font("Arial", size = 10)

pdf.multi_cell(0, 4, finalreport)

pdf_output = "finalreport.pdf"
pdf.output(pdf_output)

print(f"PDF successfully created and saved as {pdf_output}")

################################image 



def add_image_to_pdf(pdf_path, image_path, output_path):
    img = Image.open(image_path)
    orig_width, orig_height = img.size

    max_size = min(orig_width, orig_height)
    new_width = new_height = min(max_size, 150) 

    img = img.resize((new_width, new_height), Image.LANCZOS)


    packet = BytesIO()
    can = canvas.Canvas(packet, pagesize=letter)

    text_margin = 10  
    genai.configure(api_key='AIzaSyCqEKwd23ztVuk-dkCXypjeHWlcs41aCSM')
    model = genai.GenerativeModel('gemini-1.5-flash')

    image1 = {
    'mime_type': 'image/jpeg',
    'data': pathlib.Path(image_path).read_bytes()
    }

    prompt = "One title of the image?"

    response = model.generate_content([prompt, image1])
    print(response.text)

    can.drawString(100, 700, response.text)

    img_y = 700 - text_margin - new_height  

    # Draw the image on the canvas
    can.drawImage(ImageReader(img), 50, img_y, width=new_width, height=new_height)
    can.showPage()
    can.save()

    packet.seek(0)

    existing_pdf = PdfReader(open(pdf_path, "rb"))
    output = PdfWriter()

    for page in existing_pdf.pages:
        output.add_page(page)

    new_pdf = PdfReader(packet)
    output.add_page(new_pdf.pages[0])

    with open(output_path, "wb") as output_stream:
        output.write(output_stream)



image_directory = "uploads"

final_pdf_path = "finalreport.pdf"


for image_file in os.listdir(image_directory):
    if image_file.endswith(".jpg") or image_file.endswith(".jpeg") or image_file.endswith(".png"):
        image_path = os.path.join(image_directory, image_file)

        add_image_to_pdf(final_pdf_path, image_path, final_pdf_path)

print("PDF generation complete.")




file_paths = ["finalreport.pdf","data/Lists.docx"]


# Function to extract text from PDF
def extract_text_from_pdf(pdf_path):
    with open(pdf_path, "rb") as f:
        pdf_reader = PyPDF2.PdfReader(f)
        text = ""
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text()
        return text

# Function to extract text from DOCX
def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n\n'.join(full_text)

# Function to extract text from TXT
def extract_text_from_txt(txt_path):
    with open(txt_path, "r", encoding='utf-8') as f:
        return f.read()
texts = []
for path in file_paths:
    if path.endswith(".pdf"):
        texts.append(extract_text_from_pdf(path))
    elif path.endswith(".docx"):
        texts.append(extract_text_from_docx(path))
    elif path.endswith(".txt"):
        texts.append(extract_text_from_txt(path))

context = "\n\n".join(texts)

# Initialize text splitter and embeddings
text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
texts = text_splitter.split_text(context)
model = ChatGoogleGenerativeAI(model="gemini-pro",google_api_key="AIzaSyCqEKwd23ztVuk-dkCXypjeHWlcs41aCSM",
                             temperature=0.2,convert_system_message_to_human=True)
embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key="AIzaSyCqEKwd23ztVuk-dkCXypjeHWlcs41aCSM")
vector_index = Chroma.from_texts(texts, embeddings).as_retriever(search_kwargs={"k":5})

qa_chain = RetrievalQA.from_chain_type(
    model,
    retriever=vector_index,
    return_source_documents=True
)


template = """
analyse the context \n\n
List down the sections and sub sections that are metioned in the context and compare with the section needed to be inspected and list the missing \n
Example:
user:
* Engine Oil Condition: Good
    * Engine Oil Color: Clean
    * Oil Leak on Engine: No

    Sections Needed to be Inspected:
    *Engine:*
    * Engine Oil Color
    * Oil leak on engine

system:
None
{context}
Question: {question}
Helpful Answer:"""
QA_CHAIN_PROMPT = PromptTemplate.from_template(template)
qa_chain = RetrievalQA.from_chain_type(
    model,
    retriever=vector_index,
    return_source_documents=True,
    chain_type_kwargs={"prompt": QA_CHAIN_PROMPT}
)

question = """
List down only the missing inspectable items
"""
result = qa_chain({"query": question})
ans=str(result["result"])
print(ans)

# ######################   DB #####################
# from flask import Flask, request, send_file, jsonify
# import os
# import fitz  
# from werkzeug.utils import secure_filename
# import mysql.connector 
# from mysql.connector import Error


# # MySQL Configuration
# mysql_config = {
#     'host': 'localhost',
#     'database': 'caterpillar',
#     'user': 'root',
#     'password': '12Athmikha@'
# }

# def connect_to_database():
#     try:
#         connection = mysql.connector.connect(**mysql_config)
#         if connection.is_connected():
#             return connection
#     except Error as e:
#         print(f"Error connecting to MySQL: {e}")
#     return None


# def save_pdf_to_database(pdf_path):
#     connection = connect_to_database()
#     if connection is None:
#         return None
    
#     cursor = connection.cursor()
#     with open(pdf_path, 'rb') as pdf_file:
#         binary_pdf = pdf_file.read()
    
#     try:
#         cursor.execute("INSERT INTO pdfs (pdf_data) VALUES (%s)", (binary_pdf,))
#         connection.commit()
#         pdf_id = cursor.lastrowid
#     except Error as e:
#         print(f"Error inserting PDF into MySQL: {e}")
#         connection.rollback()
#         pdf_id = None
#     finally:
#         cursor.close()
#         connection.close()
    
#     return pdf_id
# save_pdf_to_database("D:/caterpillar_hackathon/finalreport.pdf")

# def retrieve_pdf_from_database(pdf_id):
#     connection = connect_to_database()
#     if connection is None:
#         return None
    
#     cursor = connection.cursor()
#     try:
#         cursor.execute("SELECT pdf_data FROM pdfs WHERE id = %s", (pdf_id,))
#         result = cursor.fetchone()
#         if result:
#             pdf_data = result[0]
#             pdf_path = os.path.join("D:/caterpillar_hackathon/statics/uploads", f"retrieved_{pdf_id}.pdf")
#             with open(pdf_path, 'wb') as pdf_file:
#                 pdf_file.write(pdf_data)
#             return pdf_path
#         else:
#             return None
#     except Error as e:
#         print(f"Error retrieving PDF from MySQL: {e}")
#         return None
#     finally:
#         cursor.close()
#         connection.close()

# retrieve_pdf_from_database(1)
